"""
Document service — ingestion & lifecycle business logic.

Pipeline for an upload:
    raw bytes ─▶ extract text (PDF/DOCX/TXT) ─▶ clean text ─▶ save file to disk
              ─▶ persist Document row (metadata + cleaned text)

Also owns listing (upload history), single-document retrieval, and deletion
(which removes both the DB row and the stored file). Routes call these
functions; they never parse files or touch the filesystem directly.
"""

from __future__ import annotations

import io
from pathlib import Path

from sqlalchemy import select
from sqlalchemy.orm import Session

from app.core.logging_config import get_logger
from app.db.models import Document
from app.services import storage
from app.services.text_cleaning import clean_text

logger = get_logger(__name__)

# Map a normalised file extension -> human label. Drives supported-type checks.
SUPPORTED_EXTENSIONS = {".pdf": "PDF", ".docx": "DOCX", ".txt": "TXT"}


class UnsupportedFileTypeError(Exception):
    """Raised when an uploaded file is not PDF, DOCX, or TXT."""


class TextExtractionError(Exception):
    """Raised when text cannot be extracted from a file."""


class DocumentNotFoundError(Exception):
    """Raised when a document does not exist or is not owned by the user."""


# ---------------------------------------------------------------------------
# Text extraction (per format)
# ---------------------------------------------------------------------------
def _resolve_extension(filename: str, content_type: str) -> str:
    """
    Determine the canonical extension (".pdf"/".docx"/".txt").

    Trusts the filename extension first (most reliable across browsers) and
    falls back to the MIME content type.
    """
    ext = Path(filename).suffix.lower()
    if ext in SUPPORTED_EXTENSIONS:
        return ext
    # Fallback by content type.
    mapping = {
        "application/pdf": ".pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
        "text/plain": ".txt",
    }
    return mapping.get(content_type, ext)


def extract_text(filename: str, content_type: str, raw: bytes) -> str:
    """
    Extract raw (uncleaned) text from file bytes based on its type.

    Raises UnsupportedFileTypeError for unknown types and TextExtractionError
    when a supported file cannot be parsed.
    """
    ext = _resolve_extension(filename, content_type)
    if ext == ".pdf":
        return _extract_pdf(raw)
    if ext == ".docx":
        return _extract_docx(raw)
    if ext == ".txt":
        return raw.decode("utf-8", errors="ignore")
    raise UnsupportedFileTypeError(
        f"Unsupported file type '{ext or content_type}'. Upload a PDF, DOCX, or TXT file."
    )


def _extract_pdf(raw: bytes) -> str:
    """Extract text from a PDF (pypdf imported lazily)."""
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover
        raise TextExtractionError("PDF support requires 'pypdf'.") from exc
    try:
        reader = PdfReader(io.BytesIO(raw))
        return "\n".join((page.extract_text() or "") for page in reader.pages)
    except Exception as exc:
        logger.exception("PDF parsing failed")
        raise TextExtractionError(f"Could not parse PDF: {exc}") from exc


def _extract_docx(raw: bytes) -> str:
    """Extract text from a DOCX, including table cells (python-docx, lazy import)."""
    try:
        from docx import Document as DocxDocument
    except ImportError as exc:  # pragma: no cover
        raise TextExtractionError("DOCX support requires 'python-docx'.") from exc
    try:
        doc = DocxDocument(io.BytesIO(raw))
        parts = [p.text for p in doc.paragraphs]
        # Include text inside tables, which paragraphs alone miss.
        for table in doc.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)
    except Exception as exc:
        logger.exception("DOCX parsing failed")
        raise TextExtractionError(f"Could not parse DOCX: {exc}") from exc


# ---------------------------------------------------------------------------
# CRUD operations
# ---------------------------------------------------------------------------
def create_document(
    db: Session, *, user_id: int, filename: str, content_type: str, raw: bytes
) -> Document:
    """
    Full ingestion: extract -> clean -> store file -> persist row.

    The raw file is saved to disk first; the DB row records its path so the two
    can be cleaned up together on delete. Raises UnsupportedFileTypeError /
    TextExtractionError on bad input.
    """
    ext = _resolve_extension(filename, content_type)
    if ext not in SUPPORTED_EXTENSIONS:
        raise UnsupportedFileTypeError(
            f"Unsupported file type '{ext or content_type}'. Upload PDF, DOCX, or TXT."
        )

    # 1) Extract + clean (do this before writing anything, so a parse failure
    #    leaves no orphaned files behind).
    raw_text = extract_text(filename, content_type, raw)
    cleaned = clean_text(raw_text)

    # 2) Persist the source file to disk.
    stored_path = storage.save_file(user_id, filename, raw)

    # 3) Record metadata + cleaned text in the database.
    document = Document(
        user_id=user_id,
        filename=filename,
        content_type=content_type,
        file_ext=ext,
        file_size=len(raw),
        storage_path=str(stored_path),
        num_chars=len(cleaned),
        num_chunks=0,  # populated by the RAG pipeline in a later week
        status="processed",
        extracted_text=cleaned,
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    logger.info(
        "Stored document | id=%s | user=%s | ext=%s | chars=%s | file=%s",
        document.id, user_id, ext, document.num_chars, filename,
    )
    return document


def list_documents_for_user(db: Session, user_id: int) -> list[Document]:
    """Return a user's documents, newest first (the upload history)."""
    stmt = (
        select(Document)
        .where(Document.user_id == user_id)
        .order_by(Document.created_at.desc())
    )
    return list(db.execute(stmt).scalars().all())


def get_document(db: Session, user_id: int, document_id: int) -> Document:
    """
    Fetch one document, enforcing ownership.

    Raises DocumentNotFoundError if it doesn't exist or belongs to another user
    (returning 404 rather than 403 avoids confirming the id exists at all).
    """
    document = db.get(Document, document_id)
    if document is None or document.user_id != user_id:
        raise DocumentNotFoundError(f"Document {document_id} not found")
    return document


def delete_document(db: Session, user_id: int, document_id: int) -> None:
    """
    Delete a document the user owns: removes the stored file then the DB row.

    Raises DocumentNotFoundError if it doesn't exist or isn't owned by the user.
    """
    document = get_document(db, user_id, document_id)
    storage.delete_file(document.storage_path)
    db.delete(document)
    db.commit()
    logger.info("Deleted document | id=%s | user=%s", document_id, user_id)
